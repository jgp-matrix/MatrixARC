#!/usr/bin/env python3
"""
Convert PDF/DOCX documents to page PNG images for OCR.
Creates a directory of page_NNN.png files for each document.
"""

import fitz  # PyMuPDF
import os
import sys

DOCS = [
    {
        "file": "CSA C22.2 No. 286 Standard 2023.pdf",
        "output_dir": "csa_c22_286_pages",
    },
    {
        "file": "UL508Asupplement 26Aug24.pdf",
        "output_dir": "ul508a_supplement_pages",
    },
    {
        "file": "UL508A 3rd Edition 26June2025.pdf",
        "output_dir": "ul508a_certification_pages",
    },
    {
        "file": "508A-NITW-FUII Follow Up and Insp Instruction .docx",
        "output_dir": "fuii_pages",
        "is_docx": True,
    },
]

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DPI = 200  # Good balance of quality vs file size


def convert_pdf(filepath, output_dir):
    """Convert PDF to page PNGs using PyMuPDF."""
    os.makedirs(output_dir, exist_ok=True)
    doc = fitz.open(filepath)
    total = len(doc)
    print(f"  {total} pages")

    for i, page in enumerate(doc):
        page_num = i + 1
        outpath = os.path.join(output_dir, f"page_{page_num:03d}.png")
        if os.path.exists(outpath):
            continue
        # Render at DPI
        mat = fitz.Matrix(DPI / 72, DPI / 72)
        pix = page.get_pixmap(matrix=mat)
        pix.save(outpath)
        if page_num % 20 == 0 or page_num == total:
            print(f"    Converted {page_num}/{total}")

    doc.close()
    print(f"  OK Done: {total} pages in {output_dir}")
    return total


def convert_docx(filepath, output_dir):
    """
    Convert DOCX to page PNGs.
    Strategy: use python-docx to extract text, render to PDF via fitz, then to PNGs.
    Simpler: extract text content and create a simple text-based image per page.
    """
    os.makedirs(output_dir, exist_ok=True)

    try:
        from docx import Document
    except ImportError:
        print("  python-docx not installed, skipping DOCX")
        return 0

    doc = Document(filepath)

    # Extract all text with paragraph formatting
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            full_text.append(" | ".join(cells))

    text = "\n".join(full_text)

    # Split into ~3000 char pages (roughly one printed page of text)
    CHARS_PER_PAGE = 3000
    pages = []
    lines = text.split("\n")
    current_page = []
    current_len = 0

    for line in lines:
        current_page.append(line)
        current_len += len(line) + 1
        if current_len >= CHARS_PER_PAGE:
            pages.append("\n".join(current_page))
            current_page = []
            current_len = 0

    if current_page:
        pages.append("\n".join(current_page))

    print(f"  {len(pages)} text pages ({len(text)} chars total)")

    # Create a simple PDF with the text, then convert to PNGs
    # Use PyMuPDF to create a PDF from text
    pdf = fitz.open()
    fontsize = 11
    margin = 50
    page_width = 612  # Letter size
    page_height = 792

    for i, page_text in enumerate(pages):
        page = pdf.new_page(width=page_width, height=page_height)
        # Insert text with word wrapping
        rect = fitz.Rect(margin, margin, page_width - margin, page_height - margin)
        page.insert_textbox(
            rect,
            page_text,
            fontsize=fontsize,
            fontname="helv",
        )

    # Save as temp PDF then convert to PNGs
    total = len(pdf)
    for i, page in enumerate(pdf):
        page_num = i + 1
        outpath = os.path.join(output_dir, f"page_{page_num:03d}.png")
        if os.path.exists(outpath):
            continue
        mat = fitz.Matrix(DPI / 72, DPI / 72)
        pix = page.get_pixmap(matrix=mat)
        pix.save(outpath)

    pdf.close()
    print(f"  OK Done: {total} pages in {output_dir}")
    return total


def main():
    print("Converting documents to page PNGs...\n")

    for doc_info in DOCS:
        filepath = os.path.join(BASE_DIR, doc_info["file"])
        output_dir = os.path.join(BASE_DIR, doc_info["output_dir"])

        if not os.path.exists(filepath):
            print(f"SKIP: {doc_info['file']} — file not found")
            continue

        print(f"Processing: {doc_info['file']}")

        if doc_info.get("is_docx"):
            convert_docx(filepath, output_dir)
        else:
            convert_pdf(filepath, output_dir)

        print()

    print("All conversions complete!")


if __name__ == "__main__":
    main()
